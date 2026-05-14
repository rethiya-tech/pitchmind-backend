import io
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build_docx(slides: list[Any]) -> bytes:
    """Generate a Word document from slide data.

    Each slide becomes:
      - A Heading 1 for the title
      - Bullet list items for each bullet
      - A speaker notes section in italic if present
    """
    doc = Document()

    # Remove default empty paragraph
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    # Document title
    title_para = doc.add_heading("Presentation", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, slide in enumerate(slides):
        title = getattr(slide, "title", None) or ""
        bullets = list(getattr(slide, "bullets", []) or [])
        notes = getattr(slide, "speaker_notes", None) or ""

        # Slide heading
        heading = doc.add_heading(f"{i + 1}. {title}" if title else f"Slide {i + 1}", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Bullets
        for bullet in bullets:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(bullet)
            run.font.size = Pt(11)

        # Speaker notes (italic, muted)
        if notes:
            notes_para = doc.add_paragraph()
            run = notes_para.add_run(f"Notes: {notes}")
            run.font.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        # Separator between slides (except last)
        if i < len(slides) - 1:
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
